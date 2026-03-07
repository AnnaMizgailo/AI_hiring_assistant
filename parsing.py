from __future__ import annotations

import datetime as dt
import json
import logging
import os
import re
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple

import httpx
from sqlalchemy import select, update
from sqlalchemy.ext.asyncio import AsyncSession

from database import Application, Candidate, Document, Profile

logger = logging.getLogger("parsing")

STORAGE_DIR = Path(os.getenv("RESUME_STORAGE_DIR", "./storage_resumes"))
OLLAMA_URL = os.getenv("OLLAMA_URL", os.getenv("OLLAMA_BASE_URL", "http://127.0.0.1:11434/api/generate"))
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "qwen3:1.7b")
OLLAMA_TIMEOUT_SECONDS = float(os.getenv("OLLAMA_TIMEOUT_SECONDS", "600"))
OLLAMA_MAX_INPUT_CHARS = int(os.getenv("OLLAMA_MAX_INPUT_CHARS", "5000"))
OLLAMA_NUM_PREDICT = int(os.getenv("OLLAMA_NUM_PREDICT", "900"))
OLLAMA_MAX_RETRIES = int(os.getenv("OLLAMA_MAX_RETRIES", "3"))

RUS_MONTHS = {
    "янв": 1, "январ": 1,
    "фев": 2, "феврал": 2,
    "мар": 3, "март": 3,
    "апр": 4, "апрел": 4,
    "май": 5, "мая": 5,
    "июн": 6, "июн": 6,
    "июл": 7, "июл": 7,
    "авг": 8, "август": 8,
    "сен": 9, "сент": 9,
    "окт": 10, "октя": 10,
    "ноя": 11, "ноябр": 11,
    "дек": 12, "дека": 12,
}
EN_MONTHS = {
    "jan": 1, "january": 1,
    "feb": 2, "february": 2,
    "mar": 3, "march": 3,
    "apr": 4, "april": 4,
    "may": 5,
    "jun": 6, "june": 6,
    "jul": 7, "july": 7,
    "aug": 8, "august": 8,
    "sep": 9, "sept": 9, "september": 9,
    "oct": 10, "october": 10,
    "nov": 11, "november": 11,
    "dec": 12, "december": 12,
}

TECHNOLOGIES = {
    "Python", "SQL", "C", "C++", "Java", "JavaScript", "TypeScript", "Go",
    "PostgreSQL", "MySQL", "SQLite", "Pandas", "NumPy", "Machine Learning",
    "Deep Learning", "NLP", "Algorithms", "Data Structures",
}
FRAMEWORKS = {
    "Django", "Flask", "FastAPI", "PyTorch", "TensorFlow", "scikit-learn", "React",
}
TOOLS = {
    "LaTeX", "Git", "GitHub", "Docker", "Linux", "Excel", "Power BI", "Tableau",
    "Figma", "Jupyter", "VS Code",
}
HUMAN_LANGUAGES = {"English", "Russian", "German", "French", "Spanish", "Chinese"}
SKILL_ALIASES = {
    "python": "Python",
    "sql": "SQL",
    "c++": "C++",
    " c ": "C",
    "latex": "LaTeX",
    "github": "GitHub",
    "git": "Git",
    "postgres": "PostgreSQL",
    "postgresql": "PostgreSQL",
    "mysql": "MySQL",
    "numpy": "NumPy",
    "pandas": "Pandas",
    "machine learning": "Machine Learning",
    "deep learning": "Deep Learning",
    "nlp": "NLP",
    "algorithms": "Algorithms",
    "алгоритмы": "Algorithms",
    "структуры данных": "Data Structures",
}
SKILL_BLACKLIST = {
    "Fiverr", "ShareMIPT", "SQL Academy", "Stepik", "MIPT", "МФТИ",
    "Olympiad", "IOAA", "ВсОШ",
}
ENGLISH_MARKERS = [
    (re.compile(r"\bc2\b|native|fluent|свободн", re.I), "C2"),
    (re.compile(r"\bc1\b|advanced", re.I), "C1"),
    (re.compile(r"\bb2\b|upper[ -]?intermediate", re.I), "B2"),
    (re.compile(r"\bb1\b|intermediate", re.I), "B1"),
    (re.compile(r"\ba2\b|pre[ -]?intermediate|elementary", re.I), "A2"),
    (re.compile(r"\ba1\b|beginner", re.I), "A1"),
]


class ResumeParsingError(RuntimeError):
    pass


class TextExtractionError(ResumeParsingError):
    pass


@dataclass
class ExtractionResult:
    text: str
    method: str
    quality_score: float
    warnings: List[str] = field(default_factory=list)


def clean_text(text: str) -> str:
    text = text or ""
    text = text.replace("\x00", " ")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\u00a0", " ").replace("\u200b", "")
    text = text.replace("–", "-").replace("—", "-").replace("−", "-")
    text = text.replace("½", " ")
    text = re.sub(r"([A-Za-zА-Яа-я])\-\n([A-Za-zА-Яа-я])", r"\1\2", text)
    text = re.sub(r"[ \t\f\v]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def normalize_whitespace(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def text_quality_score(text: str) -> float:
    cleaned = clean_text(text)
    if not cleaned:
        return 0.0
    return min(len(cleaned) / 1200.0, 1.0)


def _guess_file_kind(file_path: str, mime_type: str) -> str:
    mime = (mime_type or "").lower()
    ext = Path(file_path).suffix.lower().strip(".")
    if "pdf" in mime or ext == "pdf":
        return "pdf"
    if "word" in mime or "officedocument.wordprocessingml" in mime or ext == "docx":
        return "docx"
    return ext or "txt"


def _extract_text_from_pdf(file_path: str) -> ExtractionResult:
    warnings: List[str] = []
    text = ""
    try:
        import fitz
        parts: List[str] = []
        with fitz.open(file_path) as doc:
            for page in doc:
                parts.append(page.get_text("text", sort=True))
        text = clean_text("\n\n".join(parts))
        method = "pymupdf"
    except Exception as exc:
        logger.warning("PyMuPDF extraction failed: %s", exc)
        try:
            import pdfplumber
            parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    parts.append(page.extract_text() or "")
            text = clean_text("\n\n".join(parts))
            method = "pdfplumber"
            warnings.append("pymupdf_failed")
        except Exception as exc2:
            raise TextExtractionError(f"PDF extraction failed: {exc2}") from exc2

    if not text:
        raise TextExtractionError("Extracted PDF text is empty")
    return ExtractionResult(text=text, method=method, quality_score=text_quality_score(text), warnings=warnings)


def _extract_text_from_docx(file_path: str) -> ExtractionResult:
    try:
        from docx import Document as DocxDocument
    except Exception as exc:
        raise TextExtractionError(f"python-docx import failed: {exc}") from exc

    doc = DocxDocument(file_path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    text = clean_text("\n".join(parts))
    if not text:
        raise TextExtractionError("Extracted DOCX text is empty")
    return ExtractionResult(text=text, method="python-docx", quality_score=text_quality_score(text))


def extract_text_details(file_path: str, mime_type: str) -> ExtractionResult:
    kind = _guess_file_kind(file_path, mime_type)
    logger.info("Extracting text kind=%s path=%s", kind, file_path)
    if kind == "pdf":
        result = _extract_text_from_pdf(file_path)
    elif kind == "docx":
        result = _extract_text_from_docx(file_path)
    else:
        text = Path(file_path).read_text(errors="ignore")
        result = ExtractionResult(text=clean_text(text), method="plain-text", quality_score=text_quality_score(text))
    logger.info("Extraction done method=%s quality=%.3f warnings=%s", result.method, result.quality_score, result.warnings)
    return result


def _extract_json_object(text: str) -> Dict[str, Any]:
    text = (text or "").strip()
    if not text:
        return {}
    try:
        obj = json.loads(text)
        return obj if isinstance(obj, dict) else {}
    except Exception:
        pass
    fenced = re.search(r"```(?:json)?\s*(\{.*\})\s*```", text, flags=re.S | re.I)
    if fenced:
        try:
            obj = json.loads(fenced.group(1))
            return obj if isinstance(obj, dict) else {}
        except Exception:
            pass
    start = text.find("{")
    end = text.rfind("}")
    if start != -1 and end != -1 and end > start:
        candidate = text[start:end + 1]
        try:
            obj = json.loads(candidate)
            return obj if isinstance(obj, dict) else {}
        except Exception:
            return {}
    return {}


def _strip_thinking(text: str) -> str:
    text = text or ""
    text = re.sub(r"<think>.*?</think>", "", text, flags=re.S | re.I)
    return text.strip()


def _trim_resume_text(raw_text: str) -> str:
    text = clean_text(raw_text)
    return text[:OLLAMA_MAX_INPUT_CHARS]


def build_parsing_prompt(raw_text: str) -> str:
    resume_text = _trim_resume_text(raw_text)
    schema = {
        "full_name": "string|null",
        "summary": "string|null",
        "desired_position": "string|null",
        "location": "string|null",
        "english_level": "A1|A2|B1|B2|C1|C2|null",
        "skills": ["string"],
        "experience": [{
            "company": "string|null",
            "position": "string|null",
            "date_start": "YYYY-MM|null",
            "date_end": "YYYY-MM|null",
            "description": "string|null",
        }],
        "education": [{"institution": "string|null", "degree": "string|null"}],
        "contacts": {"email": "string|null", "phone": "string|null", "telegram": "string|null"},
    }
    return (
        "/no_think\n"
        "Return exactly one valid JSON object. No markdown. No comments. No explanations.\n"
        "Use only facts from the resume text. If a field is missing, return null or [].\n"
        "Do not infer English level, desired position, or exact dates unless they are stated or strongly supported in the text.\n"
        "If current job is mentioned, set date_end to null.\n"
        "Keep summary short and factual.\n"
        f"Schema: {json.dumps(schema, ensure_ascii=False)}\n"
        f"Resume text:\n{resume_text}"
    )


async def _ollama_generate(prompt: str, attempt: int) -> str:
    payload = {
        "model": OLLAMA_MODEL,
        "prompt": prompt,
        "stream": False,
        "format": "json",
        "think": False,
        "options": {
            "temperature": 0.0,
            "num_predict": OLLAMA_NUM_PREDICT,
            "top_p": 0.8,
            "top_k": 20,
        },
    }
    logger.info(
        "Ollama request attempt=%s model=%s timeout=%ss prompt_chars=%s",
        attempt,
        OLLAMA_MODEL,
        OLLAMA_TIMEOUT_SECONDS,
        len(prompt),
    )
    async with httpx.AsyncClient(timeout=OLLAMA_TIMEOUT_SECONDS) as client:
        response = await client.post(OLLAMA_URL, json=payload)
        response.raise_for_status()
        body = response.json()
        raw = (body.get("response") or "").strip()
        logger.info("Ollama raw response chars=%s", len(raw))
        logger.info("Ollama raw response preview=%s", raw[:700].replace("\n", " "))
        return raw


async def call_ollama_parser(raw_text: str) -> Dict[str, Any]:
    prompt = build_parsing_prompt(raw_text)
    last_reason = "unknown"
    for attempt in range(1, OLLAMA_MAX_RETRIES + 1):
        try:
            raw_response = await _ollama_generate(prompt, attempt=attempt)
            raw_response = _strip_thinking(raw_response)
            parsed = _extract_json_object(raw_response)
            if parsed:
                logger.info("Parsed JSON keys=%s", sorted(parsed.keys()))
                return parsed
            last_reason = "empty_or_invalid_json"
            logger.warning("Ollama response was empty or invalid JSON")
        except Exception as exc:
            last_reason = str(exc)
            logger.exception("Ollama parsing attempt %s failed", attempt)
    raise ResumeParsingError(f"Ollama parsing failed: {last_reason}")


def _to_list(value: Any) -> List[Any]:
    if value is None:
        return []
    if isinstance(value, list):
        return value
    return [value]


def _normalize_skill_token(token: str) -> Optional[str]:
    token = normalize_whitespace(token)
    token = token.strip("-•,.;:()[]{}\"'")
    if not token:
        return None
    lowered = token.lower()
    if lowered in SKILL_ALIASES:
        return SKILL_ALIASES[lowered]
    if f" {lowered} " in SKILL_ALIASES:
        return SKILL_ALIASES[f" {lowered} "]
    if token in SKILL_BLACKLIST:
        return None
    if len(token) > 40:
        return None
    return token


EMAIL_RE = re.compile(r"\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b")
PHONE_RE = re.compile(r"(?:(?:\+7|8)\s*[\(\- ]?\d{3}[\)\- ]?\s*\d{3}[\- ]?\d{2}[\- ]?\d{2})")
TG_RE = re.compile(r"(?:telegram|tg|телеграм|телеграмм)?[:\s]*(@[A-Za-z0-9_]{4,})", re.I)
NAME_CLEAN_RE = re.compile(r"[^A-Za-zА-Яа-яЁё\- ]")


def _extract_contacts_from_text(text: str) -> Dict[str, Optional[str]]:
    email = None
    phone = None
    telegram = None
    m = EMAIL_RE.search(text)
    if m:
        email = m.group(0)
    m = PHONE_RE.search(text)
    if m:
        phone = normalize_whitespace(m.group(0))
    m = TG_RE.search(text)
    if m:
        telegram = m.group(1)
    return {"email": email, "phone": phone, "telegram": telegram}


def _extract_name_from_header(text: str) -> Optional[str]:
    for line in text.splitlines()[:5]:
        cleaned = NAME_CLEAN_RE.sub(" ", line).strip()
        words = [w for w in cleaned.split() if len(w) > 1]
        if 2 <= len(words) <= 4 and not EMAIL_RE.search(line) and not PHONE_RE.search(line):
            return " ".join(words)
    return None


def _normalize_full_name(name: Any, raw_text: str) -> Optional[str]:
    if isinstance(name, str) and name.strip():
        cleaned = normalize_whitespace(NAME_CLEAN_RE.sub(" ", name))
        parts = cleaned.split()
        if 2 <= len(parts) <= 4:
            return " ".join(parts)
    return _extract_name_from_header(raw_text)


def _supported_english_level(level: Any, raw_text: str) -> Optional[str]:
    if isinstance(level, str):
        up = level.upper().strip()
        if up in {"A1", "A2", "B1", "B2", "C1", "C2"}:
            for pattern, resolved in ENGLISH_MARKERS:
                if pattern.search(raw_text):
                    return resolved if resolved == up else resolved
            return None
    for pattern, resolved in ENGLISH_MARKERS:
        if pattern.search(raw_text):
            return resolved
    return None


def _contains_phrase(raw_text: str, phrase: str) -> bool:
    base = normalize_whitespace(raw_text).lower()
    needle = normalize_whitespace(phrase).lower()
    return bool(needle) and needle in base


def _normalize_desired_position(value: Any, raw_text: str) -> Optional[str]:
    if not isinstance(value, str) or not value.strip():
        return None
    value = normalize_whitespace(value)
    if _contains_phrase(raw_text, value):
        return value
    return None


def _normalize_summary(value: Any) -> Optional[str]:
    if not isinstance(value, str):
        return None
    value = normalize_whitespace(value)
    if not value:
        return None
    if len(value) > 280:
        value = value[:277].rstrip() + "..."
    return value


def _parse_month(value: str) -> Optional[int]:
    low = value.lower()
    for mapping in (RUS_MONTHS, EN_MONTHS):
        for key, month in mapping.items():
            if key in low:
                return month
    return None


def _normalize_date(value: Any) -> Optional[str]:
    if value is None:
        return None
    if isinstance(value, (int, float)):
        year = int(value)
        if 1900 <= year <= 2100:
            return f"{year:04d}-01"
        return None
    text = normalize_whitespace(str(value)).lower()
    if not text:
        return None
    if text in {"present", "current", "now", "н.в.", "настоящее время", "по настоящее время", "по н.в."}:
        return None
    m = re.match(r"^(19|20)\d{2}-(0[1-9]|1[0-2])$", text)
    if m:
        return text
    m = re.match(r"^(19|20)\d{2}$", text)
    if m:
        return f"{text}-01"
    m = re.match(r"^(0?[1-9]|1[0-2])[./-]((?:19|20)\d{2})$", text)
    if m:
        month, year = int(m.group(1)), int(m.group(2))
        return f"{year:04d}-{month:02d}"
    m = re.match(r"^((?:19|20)\d{2})[./-](0?[1-9]|1[0-2])$", text)
    if m:
        year, month = int(m.group(1)), int(m.group(2))
        return f"{year:04d}-{month:02d}"
    year_match = re.search(r"(19|20)\d{2}", text)
    if year_match:
        year = int(year_match.group(0))
        month = _parse_month(text) or 1
        return f"{year:04d}-{month:02d}"
    return None


def _normalize_experience_item(item: Any) -> Optional[Dict[str, Any]]:
    if not isinstance(item, dict):
        return None
    company = item.get("company")
    position = item.get("position")
    description = item.get("description")
    if isinstance(company, str):
        company = normalize_whitespace(company.replace("•", ""))
    if isinstance(position, str):
        position = normalize_whitespace(position.replace("•", ""))
        position = position.replace("образо-", "образовательных")
    if isinstance(description, str):
        description = normalize_whitespace(description)
    out = {
        "company": company or None,
        "position": position or None,
        "date_start": _normalize_date(item.get("date_start")),
        "date_end": _normalize_date(item.get("date_end")),
        "description": description or None,
    }
    if not any(out.values()):
        return None
    return out


def _normalize_education_item(item: Any) -> Optional[Dict[str, Any]]:
    if not isinstance(item, dict):
        return None
    institution = normalize_whitespace(str(item.get("institution"))) if item.get("institution") else None
    degree = normalize_whitespace(str(item.get("degree"))) if item.get("degree") else None
    if institution == "None":
        institution = None
    if degree == "None":
        degree = None
    if not institution and not degree:
        return None
    return {"institution": institution, "degree": degree}


def _extract_skills_from_text(raw_text: str) -> List[str]:
    low = f" {normalize_whitespace(raw_text).lower()} "
    found: List[str] = []
    for alias, canonical in SKILL_ALIASES.items():
        pattern = re.escape(alias.strip())
        if re.search(rf"(?<!\w){pattern}(?!\w)", low):
            found.append(canonical)
    return sorted(set(found))


def _normalize_skills(model_skills: Iterable[Any], raw_text: str) -> List[str]:
    out: List[str] = []
    for item in model_skills:
        if isinstance(item, str):
            for part in re.split(r"[,/;|]", item):
                skill = _normalize_skill_token(part)
                if skill:
                    out.append(skill)
    out.extend(_extract_skills_from_text(raw_text))
    filtered = []
    for skill in out:
        if skill in SKILL_BLACKLIST:
            continue
        if skill.lower().startswith("курс "):
            continue
        filtered.append(skill)
    return sorted(set(filtered))


def _split_skill_buckets(skills: List[str]) -> Tuple[List[str], List[str], List[str], List[str], List[str]]:
    technologies: List[str] = []
    frameworks: List[str] = []
    tools: List[str] = []
    languages: List[str] = []
    other: List[str] = []
    for skill in skills:
        if skill in FRAMEWORKS:
            frameworks.append(skill)
        elif skill in TOOLS:
            tools.append(skill)
        elif skill in HUMAN_LANGUAGES:
            languages.append(skill)
        elif skill in TECHNOLOGIES:
            technologies.append(skill)
        else:
            other.append(skill)
    all_skills = sorted(set(technologies + frameworks + tools + languages + other))
    return all_skills, sorted(set(technologies)), sorted(set(frameworks)), sorted(set(tools)), sorted(set(languages))


def _month_diff(start: dt.datetime, end: dt.datetime) -> int:
    return max(0, (end.year - start.year) * 12 + (end.month - start.month))


def total_experience_years(experience: List[Dict[str, Any]]) -> float:
    total_months = 0
    now = dt.datetime.now()
    for exp in experience:
        start_str = exp.get("date_start")
        if not start_str:
            continue
        try:
            start = dt.datetime.strptime(start_str, "%Y-%m")
        except Exception:
            continue
        end_str = exp.get("date_end")
        try:
            end = dt.datetime.strptime(end_str, "%Y-%m") if end_str else now
        except Exception:
            end = now
        total_months += _month_diff(start, end)
    return round(total_months / 12.0, 1)


def _merge_contacts(model_contacts: Any, raw_text: str) -> Dict[str, Any]:
    extracted = _extract_contacts_from_text(raw_text)
    model_contacts = model_contacts if isinstance(model_contacts, dict) else {}
    return {
        "email": model_contacts.get("email") or extracted.get("email"),
        "phone": model_contacts.get("phone") or extracted.get("phone"),
        "telegram": model_contacts.get("telegram") or extracted.get("telegram"),
    }


def _is_profile_meaningful(profile: Dict[str, Any]) -> bool:
    signals = 0
    for key in ("full_name", "skills", "experience", "education"):
        value = profile.get(key)
        if value not in (None, [], {}, ""):
            signals += 1
    return signals >= 2


def _build_confidence(profile: Dict[str, Any], contacts: Dict[str, Any], raw_text: str) -> Dict[str, Any]:
    fields: Dict[str, float] = {}
    for key, value in profile.items():
        if value in (None, [], {}, ""):
            fields[key] = 0.15
        elif key in {"english_level", "desired_position"}:
            fields[key] = 0.55
        elif key in {"experience", "education"}:
            fields[key] = 0.8
        else:
            fields[key] = 0.9
    return {
        "fields": fields,
        "contacts": contacts,
        "extraction": {
            "method": "ollama-llm+postprocess",
            "model": OLLAMA_MODEL,
            "raw_text_length": len(raw_text),
        },
    }


def _build_missing_fields(profile: Dict[str, Any], raw_text: str) -> Dict[str, Any]:
    required_for_scoring = []
    if not profile.get("skills"):
        required_for_scoring.append("skills")
    if not profile.get("experience"):
        required_for_scoring.append("experience")
    empty_fields = [k for k, v in profile.items() if v in (None, [], {}, "")]
    return {
        "required_for_scoring": required_for_scoring,
        "empty_fields": empty_fields,
        "raw_text_length": len(raw_text),
    }


async def parse_resume_to_profile(raw_text: str) -> Tuple[Dict[str, Any], Dict[str, Any], Dict[str, Any]]:
    logger.info("Resume text chars=%s", len(raw_text))
    logger.info("Resume preview=%s", raw_text[:300].replace("\n", " "))
    llm_data = await call_ollama_parser(raw_text)

    contacts = _merge_contacts(llm_data.get("contacts"), raw_text)
    full_name = _normalize_full_name(llm_data.get("full_name"), raw_text)
    english_level = _supported_english_level(llm_data.get("english_level"), raw_text)
    desired_position = _normalize_desired_position(llm_data.get("desired_position"), raw_text)
    summary = _normalize_summary(llm_data.get("summary"))

    experience = []
    for item in _to_list(llm_data.get("experience")):
        normalized = _normalize_experience_item(item)
        if normalized:
            experience.append(normalized)

    education = []
    for item in _to_list(llm_data.get("education")):
        normalized = _normalize_education_item(item)
        if normalized:
            education.append(normalized)

    skills = _normalize_skills(_to_list(llm_data.get("skills")), raw_text)
    skills, technologies, frameworks, tools, languages = _split_skill_buckets(skills)

    profile_json = {
        "full_name": full_name,
        "summary": summary,
        "desired_position": desired_position,
        "location": normalize_whitespace(str(llm_data.get("location"))) if llm_data.get("location") else None,
        "english_level": english_level,
        "total_experience_years": total_experience_years(experience),
        "skills": skills,
        "technologies": technologies,
        "frameworks": frameworks,
        "tools": tools,
        "languages": languages,
        "experience": experience,
        "education": education,
    }

    if not _is_profile_meaningful(profile_json):
        raise ResumeParsingError("Parsed profile is empty or not meaningful")

    conf = _build_confidence(profile_json, contacts, raw_text)
    miss = _build_missing_fields(profile_json, raw_text)
    return profile_json, conf, miss


async def save_resume_file(db: AsyncSession, candidate_id: str, file_bytes: bytes, filename: str, mime_type: str) -> Dict[str, Any]:
    STORAGE_DIR.mkdir(parents=True, exist_ok=True)
    doc_id = str(uuid.uuid4())
    file_path = STORAGE_DIR / f"{doc_id}__{filename}"
    file_path.write_bytes(file_bytes)
    doc = Document(
        id=doc_id,
        candidate_id=candidate_id,
        file_name=filename,
        mime_type=mime_type,
        file_path=str(file_path),
        parse_status="PENDING",
    )
    db.add(doc)
    await db.flush()
    return {"document_id": doc_id}


async def parse_document(db: AsyncSession, document_id: str) -> Dict[str, Any]:
    result = await db.execute(select(Document).where(Document.id == document_id))
    doc = result.scalar_one_or_none()
    if not doc:
        return {"status": "ERROR", "message": "Not found"}

    try:
        ext = extract_text_details(doc.file_path, doc.mime_type)
        doc.raw_text = ext.text
        profile_json, confidence_json, missing_fields_json = await parse_resume_to_profile(ext.text)

        candidate_res = await db.execute(select(Candidate).where(Candidate.id == doc.candidate_id))
        candidate = candidate_res.scalar_one_or_none()
        if candidate:
            candidate.full_name = profile_json.get("full_name")
            candidate.contacts_json = confidence_json.get("contacts", {})

        profile_res = await db.execute(select(Profile).where(Profile.candidate_id == doc.candidate_id))
        profile = profile_res.scalar_one_or_none()
        if not profile:
            profile = Profile(candidate_id=doc.candidate_id)
            db.add(profile)

        profile.profile_json = profile_json
        profile.confidence_json = confidence_json
        profile.missing_fields_json = missing_fields_json

        doc.parse_status = "DONE"
        doc.last_error = None
        doc.parsed_at = dt.datetime.utcnow()

        await db.execute(
            update(Application)
            .where(Application.candidate_id == doc.candidate_id)
            .values(status="PROFILE_READY", updated_at=dt.datetime.utcnow())
        )
        await db.commit()
        return {
            "status": "DONE",
            "document_id": document_id,
            "parse_status": doc.parse_status,
            "profile_json": profile_json,
        }
    except Exception as exc:
        logger.exception("Parse failed for document_id=%s", document_id)
        doc.parse_status = "ERROR"
        doc.last_error = str(exc)
        doc.parsed_at = dt.datetime.utcnow()
        await db.commit()
        return {
            "status": "ERROR",
            "document_id": document_id,
            "parse_status": doc.parse_status,
            "message": str(exc),
        }
